__author__ = 'Joeri Nicolaes'
__author_email__ = 'joerinicolaes@gmail.com'

from docx import Document

class readdocument:
    """
    Read text from MS Office .docx document
    """

    def read(self, document):
        """
        Read full text from document (link)
        :param document: document to be loaded into result
        :return: utf-8 text as string
        """
        document = Document(document)
        docText = '\n\n'.join([paragraph.text.encode('utf-8') for paragraph in document.paragraphs])
        return docText

    def readParagraph(self, document):
        """
        Read full text from document (link)
        :param document: document to be loaded into result
        :return: list of utf-8 encoded paragraph strings
        """
        document = Document(document)
        result = []
        for paragraph in document.paragraphs:
            result.append(paragraph.text.encode('utf-8'))
        return result